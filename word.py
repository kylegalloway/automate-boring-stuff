import docx
import os

os.chdir('/mnt/c/Users/ckgallow/Downloads')

# d = docx.Document('demo.docx')

# para = d.paragraphs

# print(para[0].text)
# print(para[1].runs[0].text)
# print(para[1].runs[1].bold)
# para[1].runs[2].underline = True
# para[1].runs[2].italic = True
# para[1].runs[2].bold = True
# para[1].style = 'Title'

# d.save('demo-changed.docx')

# d = docx.Document()
# d.add_paragraph('Hello this is a paragraph')
# d.add_paragraph('This is another paragraph')
# p = d.paragraphs[0]
# p.add_run('This is a new run.')
# p.runs[1].bold = True
# d.save('demo-written.docx')


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getText('demo.docx'))
